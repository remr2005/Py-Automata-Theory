from docx import Document
from collections import defaultdict

def save_to_docx(data, transitions, binMatrix, filename="output.docx", intro_text="gg\n"):
    doc = Document()

    # Вводный текст
    if intro_text:
        doc.add_paragraph(intro_text)
        doc.add_paragraph("")  # Отступ после вступления

        # Таблица бинарной матрицы (симметричная по парам)
    doc.add_heading('Бинарная матрица (binMatrix)', level=1)

    # Получаем отсортированный список всех состояний
    states = sorted({s for pair in binMatrix.keys() for s in pair}, key=int)

    # Создаем таблицу (добавим заголовки: первая строка и первый столбец)
    bin_table = doc.add_table(rows=len(states) + 1, cols=len(states) + 1)
    bin_table.style = 'Table Grid'

    # Заполняем заголовки
    for idx, state in enumerate(states):
        bin_table.cell(0, idx + 1).text = state  # верхняя строка
        bin_table.cell(idx + 1, 0).text = state  # первый столбец

    # Заполняем значения
    for i, s1 in enumerate(states):
        for j, s2 in enumerate(states):
            if i == j:
                bin_table.cell(i + 1, j + 1).text = "-"  # или "—"
            else:
                key = (s1, s2) if (s1, s2) in binMatrix else (s2, s1)
                value = binMatrix.get(key, "-")
                bin_table.cell(i + 1, j + 1).text = str(value)


    # Таблица замкнутости
    doc.add_heading('Таблица замкнутости', level=1)
    headers = ["СОСТОЯНИЯ", "ПОКРЫТИЯ", "A-перех", "B-перех", "a-реак", "b-реак", "a->перех в групп", "b->перех в групп"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    for row_data in data:
        group, elements, set1, set2, a_val, b_val = row_data
        tr_a = transitions[group]["a"]
        tr_b = transitions[group]["b"]

        # Обработка реакций
        a_out = "-" if isinstance(a_val, set) and len(a_val) == 0 else list(a_val)[0] if isinstance(a_val, set) else a_val
        b_out = "-" if isinstance(b_val, set) and len(b_val) == 0 else list(b_val)[0] if isinstance(b_val, set) else b_val

        cells = [
            group,
            ", ".join(elements),
            ", ".join(sorted(set1)),
            ", ".join(sorted(set2)),
            a_out,
            b_out,
            tr_a[0],
            tr_b[0]
        ]

        row = table.add_row().cells
        for i, val in enumerate(cells):
            row[i].text = str(val)

    doc.add_paragraph("\n")

    # Таблица переходов после покрытия
    doc.add_heading('Таблица автомата после максимального покрытия', level=2)
    t_headers = ["Состояние", "a", "b"]
    t_table = doc.add_table(rows=1, cols=3)
    t_table.style = 'Table Grid'
    t_hdr_cells = t_table.rows[0].cells
    for i, h in enumerate(t_headers):
        t_hdr_cells[i].text = h

    for group, mapping in transitions.items():
        row = t_table.add_row().cells
        row[0].text = group

        a_trans = mapping.get('a', [])
        b_trans = mapping.get('b', [])

        row[1].text = (
            f"{a_trans[0]},{a_trans[1]}" if len(a_trans) == 2 else
            f"{a_trans[0]}" if len(a_trans) == 1 else
            "-"
        )
        row[2].text = (
            f"{b_trans[0]},{b_trans[1]}" if len(b_trans) == 2 else
            f"{b_trans[0]}" if len(b_trans) == 1 else
            "-"
        )

    doc.save(f"docs/{filename}")
